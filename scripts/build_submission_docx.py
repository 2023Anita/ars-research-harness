from __future__ import annotations

import argparse
import csv
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_PACKAGE_DIR = REPO_ROOT / "examples" / "nhanes-undiagnosed-diabetes" / "submission_package"
CITATION_RE = re.compile(r"\[([0-9,\-\s]+)\]")
REFERENCE_RE = re.compile(r"^(\d+)\.\s+(.*)")


@dataclass(frozen=True)
class CitationOrderedLines:
    lines: list[str]
    citation_map: dict[int, int]
    uncited_reference_numbers: list[int]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, size: int = 10, bold: bool = False, color: str | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def add_markdown_table(doc: Document, table_lines: list[str], caption: str | None = None) -> None:
    rows: list[list[str]] = []
    for line in table_lines:
        if re.match(r"^\|\s*-+", line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, size=9, bold=True)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    if len(rows[0]) == 4 and rows[0][0].lower().startswith("characteristic"):
        widths = [2.35, 1.95, 1.95, 0.75]
    elif len(rows[0]) == 4 and rows[0][0].lower().startswith("variable"):
        widths = [2.85, 1.75, 0.7, 0.7]
    elif len(rows[0]) == 4:
        widths = [2.5, 1.9, 1.7, 0.9]
    elif len(rows[0]) == 3:
        widths = [2.6, 2.4, 2.0]
    else:
        widths = [6.8 / len(rows[0])] * len(rows[0])

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[min(col_idx, len(widths) - 1)])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if col_idx > 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                style_paragraph(paragraph, size=8, bold=(row_idx == 0), color=("FFFFFF" if row_idx == 0 else None))
            if row_idx == 0:
                set_cell_shading(cell, "1F4E79")

    doc.add_paragraph()


def collect_table(lines: list[str], start: int) -> tuple[list[str], int]:
    table_lines = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1
    return table_lines, idx


def add_supplementary_table(doc: Document, title: str, path: Path) -> None:
    if not path.exists():
        return
    doc.add_heading(title, level=2)
    lines = path.read_text(encoding="utf-8").splitlines()
    table_lines = [line.strip() for line in lines if line.strip().startswith("|")]
    add_markdown_table(doc, table_lines)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "1F4E79"),
        ("Heading 3", 11, "1F4E79"),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def expand_citation_numbers(citation: str) -> list[int]:
    numbers: list[int] = []
    for part in citation.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start_text, end_text = [piece.strip() for piece in part.split("-", 1)]
            start = int(start_text)
            end = int(end_text)
            step = 1 if start <= end else -1
            numbers.extend(range(start, end + step, step))
        else:
            numbers.append(int(part))
    return numbers


def compact_citation_numbers(numbers: list[int]) -> str:
    if not numbers:
        return ""
    ranges: list[str] = []
    start = numbers[0]
    previous = numbers[0]
    for number in numbers[1:]:
        if number == previous + 1:
            previous = number
            continue
        ranges.append(f"{start}-{previous}" if start != previous else str(start))
        start = previous = number
    ranges.append(f"{start}-{previous}" if start != previous else str(start))
    return ",".join(ranges)


def split_reference_section(lines: list[str]) -> tuple[list[str], list[str], list[str]]:
    for idx, line in enumerate(lines):
        if line.strip() == "# References":
            heading = lines[idx : idx + 1]
            return lines[:idx], heading, lines[idx + 1 :]
    return lines, [], []


def parse_reference_entries(reference_lines: list[str]) -> dict[int, str]:
    entries: dict[int, str] = {}
    for line in reference_lines:
        match = REFERENCE_RE.match(line.strip())
        if match:
            entries[int(match.group(1))] = match.group(2)
    return entries


def prepare_citation_ordered_lines(lines: list[str]) -> CitationOrderedLines:
    body_lines, reference_heading, reference_lines = split_reference_section(lines)
    reference_entries = parse_reference_entries(reference_lines)
    citation_order: list[int] = []
    seen: set[int] = set()

    for line in body_lines:
        for match in CITATION_RE.finditer(line):
            for number in expand_citation_numbers(match.group(1)):
                if number not in seen:
                    seen.add(number)
                    citation_order.append(number)

    missing = [number for number in citation_order if number not in reference_entries]
    if missing:
        missing_text = ", ".join(str(number) for number in missing)
        raise ValueError(f"Missing reference entries: {missing_text}")

    citation_map = {old_number: new_number for new_number, old_number in enumerate(citation_order, start=1)}

    def rewrite_line(line: str) -> str:
        def rewrite_match(match: re.Match[str]) -> str:
            old_numbers = expand_citation_numbers(match.group(1))
            new_numbers = sorted({citation_map[number] for number in old_numbers})
            return f"[{compact_citation_numbers(new_numbers)}]"

        return CITATION_RE.sub(rewrite_match, line)

    ordered_references = [
        f"{new_number}. {reference_entries[old_number]}"
        for old_number, new_number in sorted(citation_map.items(), key=lambda item: item[1])
    ]
    uncited = sorted(set(reference_entries) - set(citation_order))
    ordered_lines = [rewrite_line(line) for line in body_lines]
    if reference_heading:
        ordered_lines.extend(reference_heading)
        ordered_lines.append("")
        ordered_lines.extend(ordered_references)
    return CitationOrderedLines(
        lines=ordered_lines,
        citation_map=citation_map,
        uncited_reference_numbers=uncited,
    )


def write_reference_order_report(ordered: CitationOrderedLines, package_dir: Path) -> None:
    rows = [["old_reference_number", "new_reference_number", "status"]]
    for old_number, new_number in sorted(ordered.citation_map.items(), key=lambda item: item[1]):
        rows.append([str(old_number), str(new_number), "cited"])
    for old_number in ordered.uncited_reference_numbers:
        rows.append([str(old_number), "", "uncited_excluded"])
    with (package_dir / "reference_order_check.csv").open("w", newline="", encoding="utf-8") as handle:
        csv.writer(handle).writerows(rows)


def build_docx(package_dir: Path = DEFAULT_PACKAGE_DIR) -> Path:
    source_md = package_dir / "manuscript_final_generic_sci.md"
    output_docx = package_dir / "manuscript_final_with_tables_figures.docx"

    doc = Document()
    configure_document(doc)

    ordered_lines = prepare_citation_ordered_lines(source_md.read_text(encoding="utf-8").splitlines())
    write_reference_order_report(ordered_lines, package_dir)
    lines = ordered_lines.lines
    idx = 0
    pending_table_caption: str | None = None
    skip_supplementary_bullets = False

    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()

        if not line or line == "---":
            idx += 1
            continue

        if line == "# Tables and Figures":
            doc.add_page_break()
            doc.add_heading("Tables and Figures", level=1)
            idx += 1
            continue

        if line == "## Supplementary Tables":
            doc.add_page_break()
            doc.add_heading("Supplementary Tables", level=1)
            add_supplementary_table(doc, "Supplementary Table 1. BMI Alternative Model", package_dir / "tables" / "supplementary_table_bmi_model.md")
            add_supplementary_table(doc, "Supplementary Table 2. Categorical Screening Model", package_dir / "tables" / "supplementary_table_categorical_model.md")
            add_supplementary_table(doc, "Supplementary Table 3. Fasting-Weighted HbA1c/FPG Sensitivity Prevalence", package_dir / "tables" / "supplementary_table_fasting_sensitivity.md")
            add_supplementary_table(doc, "Supplementary Table 4. Covariate Missingness", package_dir / "tables" / "covariate_missingness.md")
            skip_supplementary_bullets = True
            idx += 1
            continue

        if skip_supplementary_bullets and line.startswith("- Supplementary Table"):
            idx += 1
            continue
        if skip_supplementary_bullets and line.startswith("# References"):
            skip_supplementary_bullets = False

        if line.startswith("# "):
            title = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.name = "Times New Roman"
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string("1F4E79")
            idx += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
            if heading in {"Table 1", "Table 2"}:
                pending_table_caption = heading
            idx += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            idx += 1
            continue

        if line.startswith("|"):
            table_lines, idx = collect_table(lines, idx)
            add_markdown_table(doc, table_lines, pending_table_caption)
            pending_table_caption = None
            continue

        if "File: `figures/figure1_flow.png`" in line:
            p = doc.add_paragraph("Figure 1. Study population selection.")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, size=9, bold=True)
            doc.add_picture(str(package_dir / "figures" / "figure1_flow.png"), width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            idx += 1
            continue

        if "File: `figures/figure2_subgroup_prevalence.png`" in line:
            p = doc.add_paragraph("Figure 2. Weighted prevalence of HbA1c-defined undiagnosed diabetes by subgroup.")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, size=9, bold=True)
            doc.add_picture(str(package_dir / "figures" / "figure2_subgroup_prevalence.png"), width=Inches(6.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            idx += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            style_paragraph(p, size=10)
            idx += 1
            continue

        text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        text = text.replace("`", "")
        p = doc.add_paragraph(text)
        style_paragraph(p, size=10)
        idx += 1

    doc.save(output_docx)
    return output_docx


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build a DOCX manuscript with embedded tables and figures.")
    parser.add_argument("--package-dir", type=Path, default=DEFAULT_PACKAGE_DIR)
    args = parser.parse_args()
    print(build_docx(args.package_dir.resolve()))
